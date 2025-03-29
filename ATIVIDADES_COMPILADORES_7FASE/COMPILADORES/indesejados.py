import tkinter as tk
from tkinter import filedialog
from docx import Document
import pdfplumber
from fpdf import FPDF

# Função para extrair texto de arquivo Word (.docx)
def extrair_texto_word(caminho_arquivo):
    doc = Document(caminho_arquivo)
    texto = ''
    for para in doc.paragraphs:
        texto += para.text + '\n'
    return texto

# Função para extrair texto de arquivo PDF
def extrair_texto_pdf(caminho_arquivo):
    texto = ''
    with pdfplumber.open(caminho_arquivo) as pdf:
        for pagina in pdf.pages:
            texto += pagina.extract_text()
    return texto

# Função para extrair caracteres ASCII de um texto
def extrair_caracteres_ascii(texto, indesejados):
    # Filtra os caracteres indesejados
    return [(char, ord(char)) for char in texto if char not in indesejados]

# Função para abrir a caixa de diálogo e escolher o arquivo
def abrir_arquivo():
    caminho_arquivo = filedialog.askopenfilename(filetypes=[("Arquivos Word", "*.docx"), ("Arquivos PDF", "*.pdf")])
    if caminho_arquivo:
        if caminho_arquivo.endswith(".docx"):
            texto = extrair_texto_word(caminho_arquivo)
        elif caminho_arquivo.endswith(".pdf"):
            texto = extrair_texto_pdf(caminho_arquivo)
        
        # Exibir o texto extraído (primeiros 500 caracteres)
        texto_label.config(text=f"Texto extraído: {texto[:500]}...")  # Exibe os primeiros 500 caracteres

        # Extrair os caracteres ASCII
        caracteres_ascii = extrair_caracteres_ascii(texto, indesejados)
        
        # Exibir os caracteres ASCII (primeiros 20 pares char → ASCII)
        resultado_label.config(text="Caracteres ASCII extraídos:")
        resultado_texto.delete(1.0, tk.END)
        
        for char, ascii_val in caracteres_ascii[:20]:  # Exibe os primeiros 20 caracteres
            resultado_texto.insert(tk.END, f"{char} → {ascii_val}\n")

        # Atualiza a interface para mostrar mais caracteres
        mais_btn.config(state=tk.NORMAL)  # Habilita o botão para carregar mais caracteres

        # Habilitar a opção de salvar os caracteres em formato Word ou PDF
        salvar_btn.config(state=tk.NORMAL, command=lambda: salvar_arquivo(caracteres_ascii))

# Função para carregar mais caracteres ASCII
def carregar_mais():
    caminho_arquivo = filedialog.askopenfilename(filetypes=[("Arquivos Word", "*.docx"), ("Arquivos PDF", "*.pdf")])
    if caminho_arquivo:
        if caminho_arquivo.endswith(".docx"):
            texto = extrair_texto_word(caminho_arquivo)
        elif caminho_arquivo.endswith(".pdf"):
            texto = extrair_texto_pdf(caminho_arquivo)
        
        # Extrair os caracteres ASCII
        caracteres_ascii = extrair_caracteres_ascii(texto, indesejados)
        
        # Mostrar os caracteres adicionais
        for char, ascii_val in caracteres_ascii[20:40]:  # Exibe mais 20 caracteres
            resultado_texto.insert(tk.END, f"{char} → {ascii_val}\n")
        
        # Atualiza para o próximo botão "Carregar mais" até o final
        if len(caracteres_ascii) > 40:
            mais_btn.config(state=tk.NORMAL)  # Habilita o botão para carregar mais caracteres
        else:
            mais_btn.config(state=tk.DISABLED)  # Desabilita o botão se não houver mais caracteres

# Função para salvar os caracteres ASCII em um arquivo Word ou PDF
def salvar_arquivo(caracteres_ascii):
    # Escolher formato de saída (Word ou PDF)
    formato = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx"), ("PDF files", "*.pdf")])
    if formato:
        if formato.endswith(".docx"):
            salvar_word(caracteres_ascii, formato)
        elif formato.endswith(".pdf"):
            salvar_pdf(caracteres_ascii, formato)

# Função para salvar os caracteres ASCII em um arquivo Word (.docx)
def salvar_word(caracteres_ascii, caminho_arquivo):
    doc = Document()
    doc.add_heading('Caracteres ASCII extraídos', 0)
    for char, ascii_val in caracteres_ascii:
        doc.add_paragraph(f"{char} → {ascii_val}")
    doc.save(caminho_arquivo)
    print(f"Arquivo Word salvo em: {caminho_arquivo}")

# Função para salvar os caracteres ASCII em um arquivo PDF (.pdf)
def salvar_pdf(caracteres_ascii, caminho_arquivo):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Caracteres ASCII extraídos", ln=True, align="C")
    
    for char, ascii_val in caracteres_ascii:
        pdf.cell(200, 10, txt=f"{char} → {ascii_val}", ln=True, align="L")
    
    pdf.output(caminho_arquivo)
    print(f"Arquivo PDF salvo em: {caminho_arquivo}")

# Função para adicionar teclas clicadas à lista de caracteres indesejados
def adicionar_indesejado(tecla, texto_label, resultado_texto):
    if tecla not in indesejados:
        indesejados.append(tecla)
        texto_label.config(text=f"Texto extraído: Texto modificado!")
        resultado_texto.insert(tk.END, f"Removido: {tecla}\n")
        resultado_texto.tag_add("highlight", "1.0", "end")
        resultado_texto.tag_config("highlight", background="lightyellow")
    else:
        indesejados.remove(tecla)
        texto_label.config(text=f"Texto extraído: Texto restaurado!")
        resultado_texto.insert(tk.END, f"Restaurado: {tecla}\n")
        resultado_texto.tag_remove("highlight", "1.0", "end")

# Função para criar o layout do teclado
def criar_teclado():
    teclas = [
        ['q', 'w', 'e', 'r', 't', 'y', 'u', 'i', 'o', 'p'],
        ['a', 's', 'd', 'f', 'g', 'h', 'j', 'k', 'l'],
        ['z', 'x', 'c', 'v', 'b', 'n', 'm'],
        ['1', '2', '3', '4', '5', '6', '7', '8', '9', '0'],
        [' ', 'Enter', 'Backspace']
    ]
    for i, linha in enumerate(teclas):
        for j, tecla in enumerate(linha):
            btn = tk.Button(teclado_frame, text=tecla, width=5, height=2,
                            command=lambda tecla=tecla: adicionar_indesejado(tecla, texto_label, resultado_texto))
            btn.grid(row=i, column=j, padx=2, pady=2)

# Configuração da interface gráfica com Tkinter
root = tk.Tk()
root.title("Extrator de Caracteres ASCII")

# Definir o tamanho da janela
root.geometry("800x600")

# Lista para armazenar os caracteres indesejados
indesejados = []

# Frame para o teclado
teclado_frame = tk.Frame(root)
teclado_frame.pack(pady=20)

# Label de resultado (Texto extraído)
texto_label = tk.Label(root, text="Clique para abrir um arquivo Word ou PDF", wraplength=500)
texto_label.pack(pady=20)

# Label de resultado (Caracteres ASCII)
resultado_label = tk.Label(root, text="Caracteres ASCII extraídos:", wraplength=500)
resultado_label.pack(pady=20)

# Área de texto para exibir os resultados
resultado_texto = tk.Text(root, width=50, height=10)
resultado_texto.pack(pady=10)

# Botão para abrir o arquivo
abrir_btn = tk.Button(root, text="Abrir Arquivo", command=abrir_arquivo)
abrir_btn.pack(pady=10)

# Botão para carregar mais caracteres
mais_btn = tk.Button(root, text="Carregar Mais Caracteres", command=carregar_mais, state=tk.DISABLED)
mais_btn.pack(pady=10)

# Botão para salvar os caracteres em um arquivo
salvar_btn = tk.Button(root, text="Salvar Caracteres ASCII", state=tk.DISABLED)
salvar_btn.pack(pady=10)

# Criar o layout do teclado
criar_teclado()

# Executar o loop da interface gráfica
root.mainloop()
